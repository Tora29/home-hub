#!/usr/bin/env python3
"""Word設計書をMarkdownに変換する。

初回実行時に .venv を自動セットアップする（Python 3 が必要）。
見出し・段落・テーブルを Markdown に変換して出力する。

Usage:
    python extract_docx.py <path/to/file.docx>
"""
import subprocess, sys, os

VENV = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".venv")
PYTHON = os.path.join(VENV, "bin", "python3")

if sys.executable != PYTHON:
    if not os.path.exists(VENV):
        print("初回セットアップ: venv を作成中...", file=sys.stderr)
        subprocess.run([sys.executable, "-m", "venv", VENV], check=True)
    subprocess.run(
        [os.path.join(VENV, "bin", "pip"), "install", "-q", "python-docx"],
        check=True,
    )
    os.execv(PYTHON, [PYTHON] + sys.argv)

from docx import Document

HEADING = {f"Heading {i}": "#" * i for i in range(1, 7)}


def table_to_md(tbl) -> str:
    rows = []
    for i, row in enumerate(tbl.rows):
        cells = [c.text.replace("\n", " ").strip() for c in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(rows)


def main():
    if len(sys.argv) < 2:
        print("Usage: extract_docx.py <path>", file=sys.stderr)
        sys.exit(1)

    doc = Document(sys.argv[1])
    parts = []

    for block in doc.element.body:
        tag = block.tag.split("}")[-1]
        if tag == "p":
            for p in doc.paragraphs:
                if p._element is block:
                    text = p.text.strip()
                    if text:
                        prefix = HEADING.get(p.style.name, "")
                        parts.append(f"{prefix} {text}" if prefix else text)
                    break
        elif tag == "tbl":
            for t in doc.tables:
                if t._element is block:
                    parts.append(table_to_md(t))
                    break

    print("\n\n".join(parts))


main()
